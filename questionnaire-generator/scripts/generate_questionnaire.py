"""
Markdown 问卷 → Word 文档转换脚本

用法：
    python generate_questionnaire.py <input.md> [output.docx]

说明：
    读取 Markdown 格式的问卷，生成排版精美的 Word 文档。
    支持：标题层级、列表（选项）、表格（跳转逻辑表）、粗体/斜体。
"""

import re
import sys
from pathlib import Path
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("需要安装 python-docx: pip install python-docx")
    sys.exit(1)


def parse_markdown_to_docx(md_content, doc):
    """解析 Markdown 并写入 Word 文档"""
    lines = md_content.split("\n")
    i = 0
    in_table = False
    table_rows: list[list[str]] = []
    table_align = None

    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        if in_table:
            stripped = line.strip()
            if not stripped or stripped.startswith("---"):
                if stripped.startswith("|---"):
                    i += 1
                    continue
                _add_table(doc, table_rows)
                table_rows = []
                in_table = False
                i += 1
                continue
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            table_rows.append(cells)
            i += 1
            continue

        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.strip().split("|")[1:-1]]
            if cells:
                in_table = True
                table_rows.append(cells)
                i += 1
                continue

        if line.startswith("# "):
            _add_heading(doc, line[2:].strip(), 0)
        elif line.startswith("## "):
            _add_heading(doc, line[3:].strip(), 1)
        elif line.startswith("### "):
            _add_heading(doc, line[4:].strip(), 2)
        elif line.startswith("#### "):
            _add_heading(doc, line[5:].strip(), 3)
        elif line.startswith("---") or line.startswith("***"):
            doc.add_paragraph("").paragraph_format.space_before = Pt(6)
            p = doc.add_paragraph()
            run = p.add_run("─" * 50)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(180, 180, 180)
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            text = re.sub(r"^[\s]*[-*]\s+", "", line)
            _add_list_item(doc, text)
        elif re.match(r"^\d+[\.\)]\s", line):
            text = re.sub(r"^\d+[\.\)]\s+", "", line)
            _add_list_item(doc, text)
        elif line.strip().startswith("> "):
            _add_blockquote(doc, line.strip()[2:])
        elif line.strip().startswith("□ ") or line.strip().startswith("☐ ") or line.startswith("- □"):
            _add_option(doc, line.strip())
        else:
            _add_paragraph(doc, line)

        i += 1

    if in_table and table_rows:
        _add_table(doc, table_rows)


def _add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level + 1)
    for run in heading.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _add_paragraph(doc, text):
    p = doc.add_paragraph()
    _render_inline_formatting(p, text)
    p.paragraph_format.space_after = Pt(4)


def _add_list_item(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    _render_inline_formatting(p, text)
    p.paragraph_format.space_after = Pt(2)


def _add_option(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    text = text.lstrip("- ").strip()
    _render_inline_formatting(p, text)
    p.paragraph_format.space_after = Pt(1)


def _add_blockquote(doc, text):
    p = doc.add_paragraph()
    _render_inline_formatting(p, text)
    p.paragraph_format.left_indent = Cm(1)
    for run in p.runs:
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)


def _add_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text.strip())
            run.font.size = Pt(10)
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            if row_idx == 0:
                run.bold = True
    doc.add_paragraph()


def _render_inline_formatting(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(11)


def convert(md_path, docx_path):
    with open(md_path, "r", encoding="utf-8") as f:
        md_content = f.read()

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    parse_markdown_to_docx(md_content, doc)

    doc.save(docx_path)
    print(f"已生成: {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python generate_questionnaire.py <input.md> [output.docx]")
        sys.exit(1)

    md_path = Path(sys.argv[1])
    if not md_path.exists():
        print(f"文件不存在: {md_path}")
        sys.exit(1)

    if len(sys.argv) >= 3:
        docx_path = Path(sys.argv[2])
    else:
        docx_path = md_path.with_suffix(".docx")

    convert(md_path, docx_path)
