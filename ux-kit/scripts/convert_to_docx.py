#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ux-kit Markdown → DOCX 转换脚本（JSON 协议渲染）
================================================

把 ux-kit 产出的 Markdown 材料转换为 Word(.docx) 文档。

架构
----
    xxx.md → Step1 md 解析器 → 结构化 JSON blocks → Step2 python-docx 渲染 → xxx.docx

- 全文统一字体：微软雅黑（同时覆盖 ASCII / 东亚文字）
- 标题深蓝 #1F4E79，注释/引用灰 #595959
- 样式全部代码内控制，不依赖外部 reference 模板
- 支持 chart 代码块（```chart + JSON spec）用 matplotlib 渲染图表

用法
----
    python convert_to_docx.py <input.md> <output.docx>

退出码
------
    0 成功（docx 已生成）
    2 失败（调用方保留 .md 作为降级交付物）

依赖
    pip install python-docx
    图表可选：pip install matplotlib
"""

import json
import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("缺少依赖 python-docx，请先执行: pip install python-docx")

CN_FONT = "微软雅黑"
DARK = RGBColor(0x1F, 0x4E, 0x79)      # 深蓝（标题）
GRAY = RGBColor(0x59, 0x59, 0x59)      # 灰（注释/引用）


# --------------------------------------------------------------------------- #
# 字体辅助
# --------------------------------------------------------------------------- #
def _set_font(run, size=None, bold=None, italic=None, color=None):
    """为 run 设置中英文字体（微软雅黑）。"""
    run.font.name = CN_FONT
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), CN_FONT)
    rFonts.set(qn("w:hAnsi"), CN_FONT)
    rFonts.set(qn("w:eastAsia"), CN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _set_style_font(style):
    style.font.name = CN_FONT
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), CN_FONT)
    rFonts.set(qn("w:hAnsi"), CN_FONT)
    rFonts.set(qn("w:eastAsia"), CN_FONT)


def _apply_doc_fonts(doc):
    for sname in ("Normal", "Heading 1", "Heading 2", "Heading 3",
                  "List Bullet", "List Number", "Table Grid",
                  "Light Grid Accent 1", "Caption", "Title", "Subtitle"):
        try:
            _set_style_font(doc.styles[sname])
        except (KeyError, Exception):
            continue


# --------------------------------------------------------------------------- #
# 图表渲染（可选依赖 matplotlib）
# --------------------------------------------------------------------------- #
try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    _HAS_MPL = True
except Exception:
    _HAS_MPL = False


def _ensure_cjk_font():
    from matplotlib import font_manager as fm
    for name in ("Microsoft YaHei", "SimHei", "SimSun",
                 "Noto Sans CJK SC", "PingFang SC", "WenQuanYi Zen Hei"):
        try:
            fm.findfont(name, fallback_to_default=False)
            plt.rcParams["font.sans-serif"] = [name]
            break
        except Exception:
            continue
    plt.rcParams["axes.unicode_minus"] = False


def _render_chart(spec, out_png):
    if not _HAS_MPL:
        return False
    _ensure_cjk_font()
    ctype = spec.get("type")
    title = spec.get("title", "")
    figsize = tuple(spec.get("figsize", (8, 4)))

    if ctype == "bar":
        fig, ax = plt.subplots(figsize=figsize)
        labels = spec["labels"]
        series = spec["series"]
        n = len(labels)
        nser = len(series)
        width = 0.8 / nser if nser else 0.6
        for i, (sname, vals) in enumerate(series):
            xs = [j + i * width - width * (nser - 1) / 2 for j in range(n)]
            ax.bar(xs, vals, width=width, label=sname)
        ax.set_xticks(range(n))
        ax.set_xticklabels(labels, rotation=15, ha="right")
        ax.set_ylabel(spec.get("ylabel", ""))
        if nser > 1:
            ax.legend()
        ax.set_title(title)
    elif ctype == "line":
        fig, ax = plt.subplots(figsize=figsize)
        for sname, vals in spec["series"]:
            ax.plot(spec["x"], vals, marker="o", label=sname)
        ax.set_xlabel(spec.get("xlabel", ""))
        ax.set_ylabel(spec.get("ylabel", ""))
        if spec.get("series"):
            ax.legend()
        ax.set_title(title)
    elif ctype == "pie":
        fig, ax = plt.subplots(figsize=figsize)
        ax.pie(spec["values"], labels=spec["labels"], autopct="%1.1f%%",
               startangle=90)
        ax.axis("equal")
        ax.set_title(title)
    elif ctype == "scatter":
        fig, ax = plt.subplots(figsize=figsize)
        xs = [p["x"] for p in spec["points"]]
        ys = [p["y"] for p in spec["points"]]
        labels = [p.get("label", "") for p in spec["points"]]
        ax.scatter(xs, ys, s=spec.get("size", 120))
        for x, y, la in zip(xs, ys, labels):
            ax.annotate(la, (x, y), textcoords="offset points", xytext=(6, 4))
        if spec.get("quadrant_lines", True) and xs:
            ax.axhline(sum(ys) / len(ys), color=(0.6, 0.6, 0.6),
                       linestyle="--", linewidth=1)
            ax.axvline(sum(xs) / len(xs), color=(0.6, 0.6, 0.6),
                       linestyle="--", linewidth=1)
        ax.set_xlabel(spec.get("xlabel", "影响度"))
        ax.set_ylabel(spec.get("ylabel", "置信度"))
        ax.set_title(title)
    elif ctype == "funnel":
        fig, ax = plt.subplots(figsize=figsize)
        steps = spec["steps"]
        names = [s[0] for s in steps]
        vals = [s[1] for s in steps]
        ax.barh(range(len(vals)), vals)
        ax.invert_yaxis()
        ax.set_yticks(range(len(names)))
        ax.set_yticklabels(names)
        ax.set_xlabel(spec.get("xlabel", "人数"))
        ax.set_title(title)
        first = vals[0] if vals else 0
        for i, v in enumerate(vals):
            rate = "%s%%" % round(100.0 * v / first, 1) if first else ""
            ax.text(v, i, "  %s" % rate, va="center")
    elif ctype == "radar":
        fig, ax = plt.subplots(figsize=figsize, subplot_kw=dict(polar=True))
        cats = spec["categories"]
        n = len(cats)
        angles = [i / float(n) * 2 * 3.14159 for i in range(n)]
        for sname, vals in spec["series"]:
            values = vals + [vals[0]]
            angs = angles + [angles[0]]
            ax.plot(angs, values, label=sname)
            ax.fill(angs, values, alpha=0.1)
        ax.set_xticks(angles)
        ax.set_xticklabels(cats)
        if spec.get("legend", True) and spec.get("series"):
            ax.legend(loc="upper right", bbox_to_anchor=(1.3, 1.15))
        ax.set_title(title)
    else:
        raise ValueError("不支持的图表类型: %s" % ctype)

    fig.tight_layout()
    fig.savefig(out_png, dpi=150)
    plt.close(fig)
    return True


# --------------------------------------------------------------------------- #
# Step 1: Markdown → JSON blocks
# --------------------------------------------------------------------------- #
INLINE_RE = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`[^`]+`)")


def _parse_inline(text):
    """把行内 **bold** / *italic* / `code` 解析为 run 片段列表。"""
    parts = INLINE_RE.split(text)
    runs = []
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            runs.append({"text": part[2:-2], "bold": True})
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            runs.append({"text": part[1:-1], "italic": True})
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            runs.append({"text": part[1:-1], "code": True})
        else:
            runs.append({"text": part})
    return runs


def _split_row(cells):
    return [c.strip() for c in cells.strip().strip("|").split("|")]


def parse_markdown(md_path):
    """md → blocks 列表。返回 (title, blocks)。"""
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    title = None
    blocks = []
    i = 0
    n = len(lines)

    def peek_block_text(start):
        """收集连续的非空行（跳过空行），用于判断列表/表格连续性。"""
        return lines[start:]

    while i < n:
        line = lines[i].rstrip("\n")

        # 空行
        if not line.strip():
            i += 1
            continue

        # 代码块（含 chart）
        if line.strip().startswith("```"):
            lang = line.strip()[3:].strip()
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            i += 1  # skip closing ```
            code_text = "\n".join(code_lines)
            if lang == "chart":
                try:
                    spec = json.loads(code_text)
                    blocks.append({"type": "chart", "spec": spec})
                except json.JSONDecodeError:
                    blocks.append({"type": "codeblock", "text": code_text})
            else:
                blocks.append({"type": "codeblock", "text": code_text})
            continue

        # 标题
        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip("#")), 6)
            text = line.lstrip("#").strip()
            if level == 1 and title is None:
                title = text
            else:
                blocks.append({"type": "heading", "text": text, "level": min(level, 4)})
            i += 1
            continue

        # 表格
        if line.startswith("|") and i + 1 < n and re.match(r"^\s*\|[\s:\-|]+\|\s*$", lines[i + 1]):
            headers = _split_row(line)
            i += 2  # skip header + separator
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(_split_row(lines[i]))
                i += 1
            blocks.append({"type": "table", "headers": headers, "rows": rows})
            continue

        # 引用块（连续 > 行合并）
        if line.startswith(">"):
            quote_lines = []
            while i < n and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip().lstrip(">").strip())
                i += 1
            blocks.append({"type": "quote", "text": "\n".join(quote_lines)})
            continue

        # 无序列表（连续收集）
        if re.match(r"^\s*[-*]\s", line):
            items = []
            while i < n and re.match(r"^\s*[-*]\s", lines[i].rstrip("\n")):
                items.append(re.sub(r"^\s*[-*]\s", "", lines[i].strip()))
                i += 1
            blocks.append({"type": "bullets", "items": items})
            continue

        # 有序列表（连续收集）
        if re.match(r"^\s*\d+\.\s", line):
            items = []
            while i < n and re.match(r"^\s*\d+\.\s", lines[i].rstrip("\n")):
                items.append(re.sub(r"^\s*\d+\.\s", "", lines[i].strip()))
                i += 1
            blocks.append({"type": "numbered", "items": items})
            continue

        # 分隔线
        if re.match(r"^(-{3,}|\*{3,}|_{3,})$", line.strip()):
            i += 1
            continue

        # 普通段落
        blocks.append({"type": "paragraph", "text": line.strip()})
        i += 1

    return title, blocks


# --------------------------------------------------------------------------- #
# Step 2: blocks → docx 渲染
# --------------------------------------------------------------------------- #
def _add_inline_runs(paragraph, runs):
    for r in runs:
        run = paragraph.add_run(r["text"])
        _set_font(run)
        if r.get("bold"):
            run.font.bold = True
        if r.get("italic"):
            run.font.italic = True
        if r.get("code"):
            run.font.name = "Consolas"
            rpr = run._element.get_or_add_rPr()
            rFonts = rpr.get_or_add_rFonts()
            rFonts.set(qn("w:ascii"), "Consolas")
            rFonts.set(qn("w:hAnsi"), "Consolas")
            run.font.size = Pt(10)


def render_blocks(doc, title, blocks, out_docx_path):
    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    _apply_doc_fonts(doc)

    # 标题
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_inline_runs(p, _parse_inline(title))
        for run in p.runs:
            _set_font(run, size=18, bold=True, color=DARK)

    for block in blocks:
        t = block.get("type")

        if t == "heading":
            h = doc.add_heading("", level=block.get("level", 1))
            _add_inline_runs(h, _parse_inline(block["text"]))
            for run in h.runs:
                _set_font(run, bold=True, color=DARK)

        elif t == "paragraph":
            p = doc.add_paragraph()
            _add_inline_runs(p, _parse_inline(block["text"]))

        elif t == "bullets":
            for it in block["items"]:
                p = doc.add_paragraph(style="List Bullet")
                _add_inline_runs(p, _parse_inline(it))

        elif t == "numbered":
            for it in block["items"]:
                p = doc.add_paragraph(style="List Number")
                _add_inline_runs(p, _parse_inline(it))

        elif t == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            _add_inline_runs(p, _parse_inline(block["text"]))
            for run in p.runs:
                _set_font(run, size=9, color=GRAY)

        elif t == "codeblock":
            p = doc.add_paragraph()
            run = p.add_run(block["text"])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            rpr = run._element.get_or_add_rPr()
            rFonts = rpr.get_or_add_rFonts()
            rFonts.set(qn("w:ascii"), "Consolas")
            rFonts.set(qn("w:hAnsi"), "Consolas")

        elif t == "table":
            headers = block["headers"]
            rows = block["rows"]
            if not headers:
                continue
            table = doc.add_table(rows=1, cols=len(headers))
            try:
                table.style = "Light Grid Accent 1"
            except (KeyError, Exception):
                try:
                    table.style = "Table Grid"
                except Exception:
                    pass
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            for i, h in enumerate(headers):
                if i >= len(hdr):
                    break
                for para in hdr[i].paragraphs:
                    _add_inline_runs(para, _parse_inline(str(h)))
                    for run in para.runs:
                        run.font.bold = True
            for row in rows:
                cells = table.add_row().cells
                for i, val in enumerate(row):
                    if i >= len(cells):
                        break
                    for para in cells[i].paragraphs:
                        _add_inline_runs(para, _parse_inline(str(val)))
            doc.add_paragraph()

        elif t == "chart":
            if not _HAS_MPL:
                p = doc.add_paragraph("（图表生成失败：环境缺少 matplotlib）")
                _set_font(p.runs[0], color=GRAY) if p.runs else None
                continue
            png = os.path.join(os.path.dirname(os.path.abspath(out_docx_path)),
                               "_tmp_chart.png")
            try:
                if _render_chart(block["spec"], png) and os.path.exists(png):
                    doc.add_picture(png, width=Inches(6.0))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if block["spec"].get("caption"):
                        cp = doc.add_paragraph(block["spec"]["caption"])
                        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        _set_font(cp.runs[0], size=9, color=GRAY) if cp.runs else None
                    os.remove(png)
            except Exception as e:
                p = doc.add_paragraph("（图表渲染失败：%s）" % e)
                _set_font(p.runs[0], color=GRAY) if p.runs else None

        elif t == "pagebreak":
            doc.add_page_break()

        else:
            print("[warn] 忽略未知 block type: %s" % t)


# --------------------------------------------------------------------------- #
# 主流程
# --------------------------------------------------------------------------- #
def main(argv):
    if len(argv) < 3:
        print("Usage: python convert_to_docx.py <input.md> <output.docx>")
        sys.exit(1)

    md_path = argv[1]
    docx_path = argv[2]

    if not os.path.exists(md_path):
        print("ERROR: Input file not found: %s" % md_path)
        sys.exit(2)

    try:
        title, blocks = parse_markdown(md_path)
    except Exception as e:
        print("ERROR: Markdown 解析失败: %s" % e)
        sys.exit(2)

    try:
        doc = Document()
        render_blocks(doc, title, blocks, docx_path)
        doc.save(docx_path)
    except Exception as e:
        print("ERROR: DOCX 渲染失败: %s" % e)
        sys.exit(2)

    if not os.path.exists(docx_path):
        print("ERROR: DOCX 未生成")
        sys.exit(2)

    # 成功：删除中间 md
    try:
        os.remove(md_path)
    except OSError:
        pass

    print("OK: DOCX generated: %s" % docx_path)
    sys.exit(0)


if __name__ == "__main__":
    main(sys.argv)
